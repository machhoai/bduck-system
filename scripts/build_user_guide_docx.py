from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "HUONG_DAN_SU_DUNG_HE_THONG.md"
OUTPUT = ROOT / "HUONG_DAN_SU_DUNG_HE_THONG.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
GOLD = RGBColor(122, 90, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def ensure_style(document: Document, name: str, base: str | None = None):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        style.base_style = styles[base]
    return style


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_after = Pt(10)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    note = ensure_style(document, "Screenshot Note", "Normal")
    note.font.size = Pt(10)
    note.font.color.rgb = GOLD
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.right_indent = Inches(0.12)
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.15

    intro = ensure_style(document, "Manual Intro", "Normal")
    intro.font.size = Pt(11)
    intro.font.color.rgb = RGBColor(68, 68, 68)
    intro.paragraph_format.space_after = Pt(8)
    intro.paragraph_format.line_spacing = 1.25


def add_inline_markdown(paragraph, text: str, *, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold_default


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline_markdown(p, text)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline_markdown(p, text)


def add_screenshot_note(document: Document, line: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E6")
    text = line.replace("**", "")
    p = cell.paragraphs[0]
    p.style = document.styles["Screenshot Note"]
    if ":" in text:
        label, value = text.split(":", 1)
        r = p.add_run(label + ":")
        r.bold = True
        r.font.color.rgb = GOLD
        p.add_run(value)
    else:
        p.add_run(text)


def add_front_matter(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Hướng dẫn sử dụng hệ thống WMS")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Tài liệu thao tác người dùng, kèm danh sách ảnh chụp màn hình cần bổ sung")

    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660], indent_dxa=120)
    rows = [
        ("Hệ thống", "Bduck WMS"),
        ("Nguồn nội dung", "Tổng hợp từ source code frontend và bản Markdown hướng dẫn"),
        ("Ghi chú", "Các mục 'Ảnh cần chụp' là chỉ dẫn để bổ sung hình minh họa sau khi chụp giao diện thực tế."),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), "E8EEF5")
        for cell in table.rows[idx].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(10)
            table.cell(idx, 0).paragraphs[0].runs[0].bold = True

    document.add_paragraph()
    toc_heading = document.add_paragraph(style="Heading 1")
    toc_heading.add_run("Mục lục")
    toc = document.add_paragraph()
    add_toc(toc)
    document.add_page_break()


def parse_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    # Skip original H1 because front matter has a better title block.
    start = 1 if lines and lines[0].startswith("# ") else 0

    pending_number = False
    for raw in lines[start:]:
        line = raw.rstrip()
        if not line:
            pending_number = False
            continue

        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 3")
            pending_number = False
            continue
        if line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 1")
            pending_number = False
            continue
        if line.startswith("# "):
            document.add_paragraph(line[2:], style="Heading 1")
            pending_number = False
            continue

        if line.startswith("**Ảnh cần chụp") or line.startswith("**Khoanh/chú thích"):
            add_screenshot_note(document, line)
            pending_number = False
            continue

        bullet = re.match(r"^- (.+)$", line)
        if bullet:
            add_bullet(document, bullet.group(1))
            pending_number = False
            continue

        num = re.match(r"^\d+\. (.+)$", line)
        if num:
            add_number(document, num.group(1))
            pending_number = True
            continue

        if line.startswith("**") and line.endswith(":**"):
            p = document.add_paragraph()
            p.paragraph_format.keep_with_next = True
            add_inline_markdown(p, line, bold_default=True)
            pending_number = False
            continue

        p = document.add_paragraph(style="Manual Intro")
        add_inline_markdown(p, line)
        pending_number = False


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(document)

    header = section.header.paragraphs[0]
    header.text = "Hướng dẫn sử dụng hệ thống WMS"
    header.style = document.styles["Header"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    add_page_number(section.footer.paragraphs[0])
    add_front_matter(document)
    parse_markdown(document, SOURCE.read_text(encoding="utf-8"))

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
